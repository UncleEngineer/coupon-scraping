from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
import time
from bs4 import BeautifulSoup
from selenium.webdriver.chrome.options import Options





def Lazadaprice(search='iphone', page=1,sort='low',headless=False,result=True):

    '''
    sort='popularity' # best match
    sort='priceasc' # price low to high
    sort='pricedesc' # price high to low'''

    if sort=='low':
        sort='priceasc'
    elif sort=='high':
        sort='pricedesc'
    elif sort == 'best':
        sort='popularity'
    else:
        sort='priceasc'


    allresult = []
    if headless == True:
        chrome_options = Options()
        chrome_options.add_argument('--headless')  # เปิดโหมด headless
        chrome_options.add_argument('--disable-gpu')  # ปิด GPU acceleration (ช่วยลดปัญหาบางอย่าง)
        chrome_options.add_argument('--no-sandbox')  # เพิ่มความเสถียรในบางระบบ
        driver = webdriver.Chrome(options=chrome_options)
    else:
        driver = webdriver.Chrome()

    driver.get(f"https://www.lazada.co.th/catalog/?page={page}&q={search}&sort={sort}&spm=a2o4m.homepage.search.d_go")

    html = driver.page_source

    soup = BeautifulSoup(html,'html.parser')

    data = soup.find_all('div',{'class':'buTCk'})

    for i,row in enumerate(data,start=1):
        # print(row.text)
        title = row.find('div',{'class':'RfADt'})
        txt = title.a['title']
        link = 'https:' + title.a['href']
        price = row.find('div',{'class':'aBrP0'}).text
        if result == True:
            print(f'--------{i}--------')
            print(txt)
            print(price)
            print(link)
        
        price = float(price.replace('฿','').replace(',',''))
        d = [txt,price,link]
        allresult.append(d)
    
    return allresult

# headless
output = Lazadaprice('คอมพิวเตอร์',1,headless=True,result=False)

text = ''

for o in output:
    print(o)
    text += '--------------' + '\n'
    text += o[0] + '\n'
    text += f'{o[1]:,.2f}' + '\n'
    text += o[2] + '\n'

from docx import Document

document = Document()
document.add_heading('Lazada Price', 0)
p = document.add_paragraph(text)
document.save('word-output.docx')

